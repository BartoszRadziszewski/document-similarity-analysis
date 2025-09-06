"""
Document Similarity Analysis Tool
Detects and clusters similar paragraphs in documents using sentence embeddings
FIXED VERSION: Windows and macOS compatible
"""

from openai import OpenAI
import os
import sys
import docx
import fitz  # PyMuPDF for PDFs
from sklearn.metrics.pairwise import cosine_similarity, cosine_distances
from sklearn.cluster import DBSCAN
from sentence_transformers import SentenceTransformer
import numpy as np
import csv
import logging
from typing import List, Tuple, Dict, Optional
from dataclasses import dataclass
import json
from pathlib import Path
import platform

# --------------------
# CONFIGURATION - FIXED FOR CROSS-PLATFORM
# --------------------
@dataclass
class Config:
    """Configuration settings for similarity analysis"""
    api_key: Optional[str] = os.getenv("OPENAI_API_KEY")
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    chunk_size: int = 20
    eps: float = 0.3  # DBSCAN epsilon parameter
    min_samples: int = 2  # DBSCAN minimum samples
    max_summary_paragraphs: int = 10  # Max paragraphs to send for summarization
    cache_embeddings: bool = True
    log_level: str = "INFO"

config = Config()

# FIXED: Cross-platform logging setup
def setup_logging():
    """Setup cross-platform logging"""
    # Create logs directory in user's documents folder
    if platform.system() == "Windows":
        log_dir = Path.home() / "Documents" / "DocumentAnalysis" / "logs"
    elif platform.system() == "Darwin":  # macOS
        log_dir = Path.home() / "Documents" / "DocumentAnalysis" / "logs"
    else:  # Fallback for other systems
        log_dir = Path.home() / "DocumentAnalysis" / "logs"
    
    log_dir.mkdir(parents=True, exist_ok=True)
    
    logging.basicConfig(
        level=getattr(logging, config.log_level),
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_dir / "analysis.log", encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

setup_logging()
logger = logging.getLogger(__name__)

# Initialize OpenAI client with better error handling
client = None
try:
    if config.api_key:
        client = OpenAI(api_key=config.api_key)
        logger.info("✅ OpenAI client initialized")
    else:
        logger.warning("⚠️ No OpenAI API key found. Running in offline mode - summaries unavailable.")
except Exception as e:
    logger.warning(f"⚠️ OpenAI client initialization failed: {e}")

# Load embedding model with better error handling
embedder = None
try:
    embedder = SentenceTransformer(config.embedding_model)
    logger.info(f"✅ Loaded embedding model: {config.embedding_model}")
except Exception as e:
    logger.error(f"❌ Failed to load embedding model: {e}")
    sys.exit(1)

# --------------------
# FILE LOADERS - ENHANCED ERROR HANDLING
# --------------------
def load_paragraphs_from_docx(file_path: str) -> List[str]:
    """Load paragraphs from a Word document"""
    try:
        file_path = Path(file_path)  # FIXED: Use Path object
        doc = docx.Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from DOCX: {file_path.name}")
        return paragraphs
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        raise

def load_paragraphs_from_pdf(file_path: str) -> List[str]:
    """Load paragraphs from a PDF document"""
    try:
        file_path = Path(file_path)  # FIXED: Use Path object
        doc = fitz.open(file_path)
        paragraphs = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text")
            # Better paragraph detection
            page_paragraphs = []
            current_para = []
            
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    current_para.append(line)
                elif current_para:
                    paragraphs.append(" ".join(current_para))
                    current_para = []
            
            if current_para:
                paragraphs.append(" ".join(current_para))
        
        doc.close()  # FIXED: Always close the document
        logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from PDF: {file_path.name} ({len(doc)} pages)")
        return [p for p in paragraphs if len(p) > 20]  # Filter out very short paragraphs
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        raise

def load_paragraphs_from_txt(file_path: str) -> List[str]:
    """Load paragraphs from a text file"""
    try:
        file_path = Path(file_path)  # FIXED: Use Path object
        
        # FIXED: Try multiple encodings for better compatibility
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                content = file_path.read_text(encoding=encoding)
                logger.info(f"📄 Successfully read TXT file with {encoding} encoding")
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Could not decode file {file_path} with any supported encoding")
        
        # Split by double newlines for paragraph detection
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        logger.info(f"📄 Loaded {len(paragraphs)} paragraphs from TXT: {file_path.name}")
        return paragraphs
    except Exception as e:
        logger.error(f"Error loading TXT {file_path}: {e}")
        raise

def load_paragraphs(file_path: str) -> List[str]:
    """Load paragraphs from various file formats"""
    file_path = Path(file_path)  # FIXED: Use Path object consistently
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")
    
    extension = file_path.suffix.lower()
    loaders = {
        ".docx": load_paragraphs_from_docx,
        ".pdf": load_paragraphs_from_pdf,
        ".txt": load_paragraphs_from_txt
    }
    
    if extension not in loaders:
        raise ValueError(f"Unsupported file format: {extension}. Supported: {list(loaders.keys())}")
    
    return loaders[extension](str(file_path))

# --------------------
# EMBEDDINGS & CACHING - FIXED FOR CROSS-PLATFORM
# --------------------
class EmbeddingCache:
    """Cache embeddings to avoid recomputation - FIXED for cross-platform"""
    
    def __init__(self, cache_dir: Optional[str] = None):
        # FIXED: Use proper cross-platform cache directory
        if cache_dir:
            self.cache_dir = Path(cache_dir)
        else:
            # Use platform-appropriate cache location
            if platform.system() == "Windows":
                self.cache_dir = Path.home() / "Documents" / "DocumentAnalysis" / "embedding_cache"
            elif platform.system() == "Darwin":  # macOS
                self.cache_dir = Path.home() / "Documents" / "DocumentAnalysis" / "embedding_cache"
            else:  # Fallback
                self.cache_dir = Path.home() / "DocumentAnalysis" / "embedding_cache"
        
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"📦 Cache directory: {self.cache_dir}")
        
    def get_cache_path(self, text_hash: str) -> Path:
        return self.cache_dir / f"{text_hash}.npy"
    
    def get(self, texts: List[str]) -> Optional[np.ndarray]:
        """Retrieve cached embeddings if available"""
        text_hash = str(hash("".join(texts)))
        cache_path = self.get_cache_path(text_hash)
        
        try:
            if cache_path.exists():
                logger.info("📦 Loading embeddings from cache")
                return np.load(cache_path)
        except Exception as e:
            logger.warning(f"⚠️ Failed to load cache: {e}")
        
        return None
    
    def save(self, texts: List[str], embeddings: np.ndarray):
        """Save embeddings to cache"""
        try:
            text_hash = str(hash("".join(texts)))
            cache_path = self.get_cache_path(text_hash)
            np.save(cache_path, embeddings)
            logger.info("💾 Saved embeddings to cache")
        except Exception as e:
            logger.warning(f"⚠️ Failed to save cache: {e}")

# FIXED: Initialize cache with proper error handling
cache = None
if config.cache_embeddings:
    try:
        cache = EmbeddingCache()
    except Exception as e:
        logger.warning(f"⚠️ Cache initialization failed: {e}")

def get_embeddings(paragraphs: List[str]) -> np.ndarray:
    """Generate or retrieve cached embeddings for paragraphs"""
    if not paragraphs:
        return np.array([])
    
    if not embedder:
        raise RuntimeError("Embedding model not available")
    
    # Check cache
    if cache:
        cached = cache.get(paragraphs)
        if cached is not None:
            return cached
    
    # Generate new embeddings
    logger.info(f"🔄 Generating embeddings for {len(paragraphs)} paragraphs...")
    try:
        embeddings = embedder.encode(paragraphs, convert_to_tensor=False, show_progress_bar=True)
        
        # Save to cache
        if cache:
            cache.save(paragraphs, embeddings)
        
        return embeddings
    except Exception as e:
        logger.error(f"❌ Failed to generate embeddings: {e}")
        raise

# --------------------
# CLUSTERING - NO CHANGES NEEDED
# --------------------
def cluster_paragraphs(embeddings: np.ndarray, eps: float = None, min_samples: int = None) -> np.ndarray:
    """Cluster paragraphs based on cosine similarity"""
    if len(embeddings) == 0:
        return np.array([])
    
    eps = eps or config.eps
    min_samples = min_samples or config.min_samples
    
    logger.info(f"🔍 Clustering with eps={eps}, min_samples={min_samples}")
    
    # Calculate cosine distance matrix
    dist_matrix = cosine_distances(embeddings)
    
    # Apply DBSCAN clustering
    clustering = DBSCAN(eps=eps, min_samples=min_samples, metric="precomputed")
    labels = clustering.fit_predict(dist_matrix)
    
    # Log clustering results
    unique_labels = set(labels)
    n_clusters = len(unique_labels) - (1 if -1 in unique_labels else 0)
    n_noise = list(labels).count(-1)
    logger.info(f"📊 Found {n_clusters} clusters, {n_noise} unique paragraphs")
    
    return labels

def annotate_paragraphs(paragraphs: List[str], labels: np.ndarray) -> List[str]:
    """Add cluster labels to paragraphs"""
    annotated = []
    for text, label in zip(paragraphs, labels):
        if label == -1:
            annotated.append(text)
        else:
            annotated.append(f"[SIMILAR-{label:02d}] {text}")
    return annotated

# --------------------
# SUMMARIZATION - ENHANCED ERROR HANDLING
# --------------------
def detect_language(text: str) -> str:
    """Simple language detection based on character frequency"""
    polish_chars = set("ąćęłńóśźżĄĆĘŁŃÓŚŹŻ")
    polish_count = sum(1 for char in text if char in polish_chars)
    return "pl" if polish_count > len(text) * 0.01 else "en"

def summarize_group(paragraphs: List[str], label: int, model_name: str = "gpt-4o") -> str:
    """Generate AI summary for a group of similar paragraphs"""
    if not client:
        sample = paragraphs[:3] if len(paragraphs) > 3 else paragraphs
        return f"⚠️ Offline mode - no summary available. Sample paragraphs:\n" + "\n".join(sample[:200] for sample in sample)
    
    # Limit paragraphs sent to API
    sample_paragraphs = paragraphs[:config.max_summary_paragraphs]
    
    # Detect language
    combined_text = " ".join(sample_paragraphs)
    language = detect_language(combined_text)
    lang_instruction = "Respond in Polish" if language == "pl" else "Respond in English"
    
    prompt = f"""
You are analyzing a document for duplicate content. These paragraphs were identified as similar.
{lang_instruction}.

Paragraphs from cluster SIMILAR-{label:02d} ({len(paragraphs)} total):

{chr(10).join(sample_paragraphs)}

Provide a 2-4 sentence summary explaining:
1. The main topic or theme
2. Why these paragraphs are similar
3. The type of information repeated

Summary:
"""
    
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"❌ Summarization failed: {e}")
        return f"Error generating summary: {str(e)}"

# --------------------
# OUTPUT FUNCTIONS - FIXED FOR CROSS-PLATFORM
# --------------------
def save_to_docx(annotated_paragraphs: List[str], summaries: Dict[str, str], output_path: str):
    """Save results to Word document"""
    try:
        output_path = Path(output_path)  # FIXED: Use Path object
        output_path.parent.mkdir(parents=True, exist_ok=True)  # FIXED: Ensure directory exists
        
        doc = docx.Document()
        
        # Title
        doc.add_heading("Document Similarity Analysis", 0)
        
        # Summary statistics
        doc.add_heading("Summary Statistics", 1)
        doc.add_paragraph(f"Total paragraphs: {len(annotated_paragraphs)}")
        doc.add_paragraph(f"Similar clusters found: {len(summaries)}")
        unique_count = sum(1 for p in annotated_paragraphs if not p.startswith("[SIMILAR-"))
        doc.add_paragraph(f"Unique paragraphs: {unique_count}")
        
        # Annotated document
        doc.add_page_break()
        doc.add_heading("Annotated Document", 1)
        for para in annotated_paragraphs:
            p = doc.add_paragraph(para)
            if para.startswith("[SIMILAR-"):
                # Highlight similar paragraphs
                p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        
        # Group summaries
        doc.add_page_break()
        doc.add_heading("Cluster Summaries", 1)
        for group, summary in sorted(summaries.items()):
            doc.add_heading(group, level=2)
            doc.add_paragraph(summary)
        
        doc.save(output_path)
        logger.info(f"📄 Saved results to {output_path}")
        
    except Exception as e:
        logger.error(f"❌ Failed to save DOCX: {e}")
        raise

def save_to_csv(annotated_paragraphs: List[str], output_path: str):
    """Save results to CSV file"""
    try:
        output_path = Path(output_path)  # FIXED: Use Path object
        output_path.parent.mkdir(parents=True, exist_ok=True)  # FIXED: Ensure directory exists
        
        with open(output_path, mode="w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Paragraph", "Group", "Length"])
            
            for para in annotated_paragraphs:
                if para.startswith("[SIMILAR-"):
                    parts = para.split("] ", 1)
                    group = parts[0].replace("[", "")
                    text = parts[1] if len(parts) > 1 else ""
                else:
                    group = "UNIQUE"
                    text = para
                writer.writerow([text, group, len(text)])
        
        logger.info(f"📊 Saved CSV to {output_path}")
        
    except Exception as e:
        logger.error(f"❌ Failed to save CSV: {e}")
        raise

def save_to_json(annotated_paragraphs: List[str], summaries: Dict[str, str], output_path: str):
    """Save results to JSON file"""
    try:
        output_path = Path(output_path)  # FIXED: Use Path object
        output_path.parent.mkdir(parents=True, exist_ok=True)  # FIXED: Ensure directory exists
        
        data = {
            "statistics": {
                "total_paragraphs": len(annotated_paragraphs),
                "clusters": len(summaries),
                "unique": sum(1 for p in annotated_paragraphs if not p.startswith("[SIMILAR-"))
            },
            "paragraphs": [],
            "summaries": summaries
        }
        
        for para in annotated_paragraphs:
            if para.startswith("[SIMILAR-"):
                parts = para.split("] ", 1)
                group = parts[0].replace("[", "")
                text = parts[1] if len(parts) > 1 else ""
            else:
                group = "UNIQUE"
                text = para
            
            data["paragraphs"].append({
                "text": text,
                "group": group,
                "length": len(text)
            })
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"📋 Saved JSON to {output_path}")
        
    except Exception as e:
        logger.error(f"❌ Failed to save JSON: {e}")
        raise

# --------------------
# MAIN PROCESSING - ENHANCED ERROR HANDLING
# --------------------
def process_document(
    file_path: str, 
    model_name: str = "gpt-4o",
    eps: float = None,
    min_samples: int = None,
    chunk_size: int = None
) -> Tuple[List[str], Dict[str, str]]:
    """
    Process a document to find and summarize similar paragraphs
    FIXED VERSION with better error handling and cross-platform support
    """
    logger.info(f"📚 Processing document: {Path(file_path).name}")
    
    try:
        # Load paragraphs
        paragraphs = load_paragraphs(file_path)
        logger.info(f"📄 Total paragraphs: {len(paragraphs)}")
        
        if len(paragraphs) == 0:
            logger.warning("No paragraphs found in document")
            return [], {}
        
        # Get embeddings
        embeddings = get_embeddings(paragraphs)
        
        # Cluster
        labels = cluster_paragraphs(embeddings, eps, min_samples)
        
        # Annotate
        annotated = annotate_paragraphs(paragraphs, labels)
        
        # Group for summarization
        grouped = {}
        for para, label in zip(paragraphs, labels):
            if label != -1:
                grouped.setdefault(label, []).append(para)
        
        # Generate summaries
        summaries = {}
        for label, group_paragraphs in grouped.items():
            try:
                summary = summarize_group(group_paragraphs, label, model_name)
                summaries[f"SIMILAR-{label:02d}"] = summary
            except Exception as e:
                logger.warning(f"⚠️ Failed to summarize cluster {label}: {e}")
                summaries[f"SIMILAR-{label:02d}"] = f"Summary unavailable: {str(e)}"
        
        logger.info(f"✅ Processing complete: {len(summaries)} clusters found")
        return annotated, summaries
        
    except Exception as e:
        logger.error(f"❌ Document processing failed: {e}")
        raise

# FIXED: Add convenience function for getting platform-appropriate project directory
def get_project_directory(project_name: str) -> Path:
    """Get appropriate project directory for the current platform"""
    if platform.system() == "Windows":
        base_dir = Path.home() / "Documents" / "DocumentAnalysis" / "Projects"
    elif platform.system() == "Darwin":  # macOS
        base_dir = Path.home() / "Documents" / "DocumentAnalysis" / "Projects"
    else:  # Fallback
        base_dir = Path.home() / "DocumentAnalysis" / "Projects"
    
    project_dir = base_dir / project_name
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir
